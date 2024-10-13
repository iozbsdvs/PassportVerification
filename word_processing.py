import pandas as pd
import docx
import logging


def filter_ip_address_tables(word_file):
    try:
        logging.info("Начинаем фильтрацию данных из Word файла.")
        doc = docx.Document(word_file)
        relevant_sections = [
            "Система управления базами данных Pangolin SE",
            "Список виртуальных серверов Ведомства",
            "Список виртуальных серверов НСУД"
        ]

        data = []
        inside_relevant_section = False

        for para in doc.paragraphs:
            if any(section in para.text for section in relevant_sections):
                inside_relevant_section = True
                logging.info(f"Обнаружен раздел: {para.text}")

            if inside_relevant_section and para.text.strip() == "":
                inside_relevant_section = False

        # Извлечение таблиц с IP-адресами
        for table in doc.tables:
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]

            # Проверка на наличие заголовка "Имя сервера" или "Доменное имя"
            if "доменное имя" in header_cells or "имя сервера" in header_cells:
                name_col_index = header_cells.index(
                    "доменное имя") if "доменное имя" in header_cells else header_cells.index("имя сервера")
                ip_col_index = header_cells.index("ip адрес")

                logging.info("Обнаружена таблица с IP-адресами.")

                for row in table.rows[1:]:  # Пропускаем заголовки
                    if len(row.cells) > max(name_col_index, ip_col_index):
                        name = row.cells[name_col_index].text.strip()  # Имя сервера
                        ip_address = row.cells[ip_col_index].text.strip()  # IP адрес
                        sizing = "N/A"  # Если сайзинга нет в таблице, оставляем его пустым

                        if name and ip_address:
                            data.append({
                                'Имя сервера': name,
                                'Сайзинг': sizing,
                                'IP адрес': ip_address
                            })
                            logging.info(f"Добавлена ВМ: {name}, IP: {ip_address}")

        return pd.DataFrame(data)
    except Exception as e:
        logging.error(f"Ошибка при обработке Word файла: {e}")
        return pd.DataFrame()
